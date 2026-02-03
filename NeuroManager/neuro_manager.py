import os
import sys
import re
import gradio as gr
from typing import Dict, Any, List, Tuple, Optional
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from openai import OpenAI
import docx

class NeuroProjectManagerRAG:
    """
    Двухконтурная RAG-система для нейро-менеджера проектов
    """
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.vector_store = None
        self.retriever = None
        self.openai_client = None
        
        # Загрузка базы знаний из файла
        self.knowledge_text = self._load_docx_file()
        
        # Инициализация моделей
        self._init_models()
        
        # Подготовка RAG системы
        self._setup_rag_system()
    
    def _load_docx_file(self) -> str:
        """Загрузка текста из .docx файла"""
        print(f"📖 Загрузка базы знаний из: {self.docx_path}")
        
        try:
            doc = docx.Document(self.docx_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            knowledge_text = "\n".join(full_text)
            
            if not knowledge_text:
                raise ValueError("Файл .docx пуст или не содержит текста")
            
            print(f"✅ Загружено {len(knowledge_text)} символов")
            
            return knowledge_text
            
        except FileNotFoundError:
            error_msg = f"❌ Файл не найден: {self.docx_path}"
            print(error_msg)
            raise FileNotFoundError(error_msg)
        except Exception as e:
            error_msg = f"❌ Ошибка загрузки .docx файла: {str(e)}"
            print(error_msg)
            raise
    
    def _init_models(self):
        """Инициализация моделей"""
        print("🔄 Инициализация моделей...")
        
        # Основная модель через Vsegpt
        try:
            self.openai_client = OpenAI(
                api_key="sk-or-vv-8fad4ae2ddfef897cb7eb3ef2cf982a91b96de4cc7465c694785186d31a93e91",
                base_url="https://api.vsegpt.ru/v1"
            )
            self.main_model = "gpt-3.5-turbo"
            print("✅ Vsegpt API (GPT) инициализирован, модель: gpt-3.5-turbo")
        except Exception as e:
            print(f"⚠️  Ошибка инициализации Vsegpt API: {e}")
            print("   Проверьте API ключ и подключение")
            self.openai_client = None
            self.main_model = None
        
        # Эмбеддинги для векторной БД
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
        print("✅ Эмбеддинги загружены")
    
    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """Извлечение разделов из текста базы знаний"""
        print("\n🔍 Извлечение разделов из документа...")
        
        sections = {}
        lines = text.split('\n')
        
        # Определяем основные разделы по заголовкам
        section_headers = {
            "1. общая миссия и принципы": "общее",
            "2. этапы управления ml-проектом": "этапы",
            "3. решение типовых проблем ml-проектов": "проблемы", 
            "4. технологический стек (рекомендуемый)": "технологии",
            "5. шаблоны и чек-листы": "чеклисты",
            "6. интеграция с архитектурой": "архитектура"
        }
        
        current_section = None
        current_content = []
        
        for line in lines:
            line_lower = line.strip().lower()
            line_original = line.strip()
            
            # Проверяем, является ли строка заголовком раздела
            found_section = False
            for header, category in section_headers.items():
                if header in line_lower:
                    # Сохраняем предыдущий раздел
                    if current_section:
                        sections[current_section] = "\n".join(current_content)
                    
                    # Начинаем новый раздел
                    current_section = category
                    current_content = [line_original]
                    found_section = True
                    print(f"   ✅ Найден раздел: '{line_original[:50]}...' -> категория '{category}'")
                    break
            
            if not found_section and current_section:
                # Добавляем строку к текущему разделу
                if line_original:
                    current_content.append(line_original)
        
        # Сохраняем последний раздел
        if current_section:
            sections[current_section] = "\n".join(current_content)
        
        # Дополнительно разбиваем этапы на подразделы
        if "этапы" in sections:
            этапы_текст = sections["этапы"]
            подэтапы = self._extract_substages(этапы_текст)
            sections["этапы"] = этапы_текст  # Сохраняем полный текст
            
            # Сохраняем подэтапы отдельно для лучшего поиска
            for sub_name, sub_content in подэтапы.items():
                sections[f"этапы_{sub_name}"] = sub_content
        
        print(f"📊 Извлечено разделов: {len(sections)}")
        for category, content in sections.items():
            if len(content) > 100:
                print(f"   • {category}: {len(content)} символов")
        
        return sections
    
    def _extract_substages(self, text: str) -> Dict[str, str]:
        """Извлечение подэтапов из раздела этапов"""
        substages = {}
        lines = text.split('\n')
        
        current_substage = None
        current_content = []
        
        # Ищем подэтапы по шаблону "Этап X: Название"
        substage_pattern = re.compile(r'^этап\s+(\d+):\s*(.+)$', re.IGNORECASE)
        
        for line in lines:
            match = substage_pattern.match(line.strip())
            if match:
                # Сохраняем предыдущий подэтап
                if current_substage:
                    substages[current_substage] = "\n".join(current_content)
                
                # Начинаем новый подэтап
                stage_num = match.group(1)
                stage_name = match.group(2).strip().lower()
                current_substage = f"{stage_num}_{stage_name}"
                current_content = [line.strip()]
            elif current_substage and line.strip():
                current_content.append(line.strip())
        
        # Сохраняем последний подэтап
        if current_substage:
            substages[current_substage] = "\n".join(current_content)
        
        print(f"   📑 Извлечено подэтапов: {len(substages)}")
        for name in substages:
            print(f"     - {name}")
        
        return substages
    
    def _setup_rag_system(self):
        """Настройка RAG системы с улучшенным поиском"""
        print("🔧 Создание векторной базы знаний...")
        
        # Извлечение разделов
        sections = self._extract_sections_from_text(self.knowledge_text)
        
        all_docs = []
        
        # Создаем документы из извлеченных разделов
        for category, content in sections.items():
            if content and len(content) > 20:  # Пропускаем очень короткие
                # Создаем основной документ для раздела
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": self.docx_path,
                        "category": category.split('_')[0] if '_' in category else category,
                        "subcategory": category if '_' in category else None,
                        "is_primary": True,
                        "length": len(content)
                    }
                )
                all_docs.append(doc)
                
                # Если раздел большой, разбиваем на чанки
                if len(content) > 500:
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=400,
                        chunk_overlap=100,
                        separators=["\n\n", "\n", ". ", " ", ""],
                        length_function=len
                    )
                    
                    chunk_docs = splitter.split_documents([doc])
                    all_docs.extend(chunk_docs[1:])  # Первый чанк уже добавлен
                    print(f"   📄 Раздел '{category}' разбит на {len(chunk_docs)} чанков")
        
        print(f"📊 Итого документов: {len(all_docs)}")
        
        # Создаем векторное хранилище
        self.vector_store = Chroma.from_documents(
            documents=all_docs,
            embedding=self.embeddings,
            persist_directory="./chroma_neuro_db",
            collection_name="neuro_manager_knowledge"
        )
        
        # Настраиваем ретривер
        self.retriever = self.vector_store.as_retriever(
            search_type="mmr",  # Используем MMR для разнообразия результатов
            search_kwargs={
                "k": 10,  # Берем больше результатов
                "fetch_k": 20,  # Для MMR
                "lambda_mult": 0.5  # Баланс между релевантностью и разнообразием
            }
        )
        
        print("✅ RAG система настроена!")
        
        # Сохраняем все документы для ручного поиска при необходимости
        self.all_docs = all_docs
        
        # Тестируем поиск с конкретными запросами
        self._test_search_specific()
    
    def _test_search_specific(self):
        """Тестирование поиска конкретных запросов"""
        print("\n🧪 ТЕСТИРОВАНИЕ ПОИСКА ПО КОНКРЕТНЫМ ЗАПРОСАМ:")
        
        test_queries = [
            ("разбить проект на подзадачи", ["декомпозиция", "этап 1", "подзадачи"]),
            ("технологический стек", ["pytorch", "tensorflow", "mlflow"]),
            ("переобучение модели", ["переобучение", "регуляризация", "dropout"]),
            ("чек-лист сдачи", ["чек-лист", "перед сдачей", "github"]),
            ("этапы ml проекта", ["этап 0", "этап 1", "этап 2", "этап 3", "этап 4"])
        ]
        
        for query, expected_keywords in test_queries:
            print(f"\n  Запрос: '{query}'")
            print(f"  Ожидаемые ключевые слова: {expected_keywords}")
            
            try:
                docs = self.retriever.invoke(query)
                if docs:
                    print(f"    ✅ Найдено {len(docs)} документов")
                    
                    found_keywords = []
                    for i, doc in enumerate(docs[:3]):
                        content_lower = doc.page_content.lower()
                        category = doc.metadata.get("category", "неизвестно")
                        subcategory = doc.metadata.get("subcategory", "")
                        
                        # Проверяем наличие ключевых слов
                        doc_keywords = []
                        for keyword in expected_keywords:
                            if keyword in content_lower:
                                doc_keywords.append(keyword)
                                found_keywords.append(keyword)
                        
                        preview = doc.page_content[:120].replace('\n', ' ')
                        print(f"      [{i+1}] Категория: {category}{'/' + subcategory if subcategory else ''}")
                        print(f"          Найдено ключевых слов: {doc_keywords}")
                        print(f"          Предпросмотр: {preview}...")
                    
                    # Проверяем, какие ключевые слова найдены
                    missing = [kw for kw in expected_keywords if kw not in found_keywords]
                    if missing:
                        print(f"      ⚠️ Не найдены ключевые слова: {missing}")
                    else:
                        print(f"      ✅ Все ключевые слова найдены!")
                else:
                    print(f"    ❌ Документы не найдены")
                    
            except Exception as e:
                print(f"    ⚠️ Ошибка: {e}")
    
    def _find_relevant_docs_manual(self, query: str) -> List[Document]:
        """Ручной поиск релевантных документов (резервный метод)"""
        query_lower = query.lower()
        
        # Специальные правила для частых запросов
        query_to_keywords = {
            "разбить проект на подзадачи": ["декомпозиция", "этап 1", "подзадачи", "разбиение"],
            "разделить проект на части": ["декомпозиция", "этап 1", "подзадачи"],
            "как разбить проект": ["декомпозиция", "этап 1", "разбить"],
            "декомпозиция проекта": ["декомпозиция", "этап 1"],
            "подзадачи проекта": ["декомпозиция", "подзадачи", "этап 1"]
        }
        
        # Ищем по ключевым словам
        keywords = []
        for q_pattern, kw_list in query_to_keywords.items():
            if q_pattern in query_lower:
                keywords.extend(kw_list)
                break
        
        # Если не нашли специальных правил, используем общие ключевые слова
        if not keywords:
            general_keywords = {
                "этап": ["этап"],
                "подзадач": ["декомпозиция", "подзадачи"],
                "разбить": ["декомпозиция", "разбить"],
                "разделить": ["декомпозиция", "разделить"]
            }
            
            for word, kw_list in general_keywords.items():
                if word in query_lower:
                    keywords.extend(kw_list)
        
        # Ищем документы с этими ключевыми словами
        relevant_docs = []
        for doc in self.all_docs:
            content_lower = doc.page_content.lower()
            doc_keywords = []
            
            for keyword in keywords:
                if keyword in content_lower:
                    doc_keywords.append(keyword)
            
            if doc_keywords:
                # Оцениваем релевантность по количеству найденных ключевых слов
                relevance_score = len(doc_keywords)
                relevant_docs.append((relevance_score, doc))
        
        # Сортируем по релевантности
        relevant_docs.sort(key=lambda x: x[0], reverse=True)
        
        # Берем топ-5
        return [doc for score, doc in relevant_docs[:5]]
    
    def _retrieve_context(self, query: str, category: str) -> Tuple[str, List[Document]]:
        """Поиск контекста в базе знаний"""
        print(f"[SEARCH] Запрос: '{query}', Категория: '{category}'")
        
        try:
            # Сначала пробуем стандартный поиск
            docs = self.retriever.invoke(query)
            
            if not docs:
                print(f"[SEARCH] ❌ Стандартный поиск не дал результатов")
                # Пробуем ручной поиск
                docs = self._find_relevant_docs_manual(query)
                if docs:
                    print(f"[SEARCH] 🔍 Ручной поиск нашел {len(docs)} документов")
                else:
                    print(f"[SEARCH] ❌ Не найдено документов по запросу")
                    return "", []
            
            print(f"[SEARCH] ✅ Найдено документов: {len(docs)}")
            
            # Фильтруем по категории (если указана)
            relevant_docs = []
            if category and category != "другое":
                for doc in docs:
                    doc_category = doc.metadata.get("category", "")
                    if doc_category == category:
                        relevant_docs.append(doc)
                
                if not relevant_docs:
                    print(f"[SEARCH] ⚠️ Не найдено документов категории '{category}', беру все найденные")
                    relevant_docs = docs[:5]
            else:
                relevant_docs = docs[:5]
            
            # Формируем контекст
            context_parts = []
            for i, doc in enumerate(relevant_docs[:5]):
                content = doc.page_content
                doc_category = doc.metadata.get("category", "неизвестно")
                subcategory = doc.metadata.get("subcategory", "")
                
                # Если документ слишком длинный, обрезаем
                if len(content) > 600:
                    content = content[:550] + "..."
                
                context_parts.append(f"【{doc_category}{'/' + subcategory if subcategory else ''}】\n{content}")
                
                print(f"[SEARCH]   📄 Документ {i+1}:")
                print(f"            Категория: {doc_category}{'/' + subcategory if subcategory else ''}")
                print(f"            Длина: {len(doc.page_content)} символов")
                if i == 0:
                    preview = content[:100].replace('\n', ' ')
                    print(f"            Предпросмотр: {preview}...")
            
            context = "\n\n" + "="*60 + "\n\n".join(context_parts) + "\n" + "="*60
            
            return context, relevant_docs
            
        except Exception as e:
            print(f"[SEARCH] ⚠️ Ошибка поиска: {str(e)[:100]}")
            return "", []
    
    def _generate_strict_response(self, query: str, context: str, category: str, docs: List[Document]) -> str:
        """Генерация ответа строго по документам"""
        
        # Если нет контекста, сообщаем об отсутствии информации
        if not context or not docs:
            return self._no_info_response(query)
        
        # Для запросов об этапах - особая проверка
        query_lower = query.lower()
        if any(word in query_lower for word in ["этап", "фаза", "стадия"]):
            # Проверяем, есть ли в документах информация об этапах
            has_stage_info = False
            for doc in docs:
                content_lower = doc.page_content.lower()
                if any(stage_word in content_lower for stage_word in ["этап 0", "этап 1", "этап 2", "этап 3", "этап 4"]):
                    has_stage_info = True
                    break
            
            if not has_stage_info:
                # Проверяем другие документы
                all_stage_docs = []
                for doc in self.all_docs:
                    content_lower = doc.page_content.lower()
                    if any(stage_word in content_lower for stage_word in ["этап 0", "этап 1", "этап 2", "этап 3", "этап 4"]):
                        all_stage_docs.append(doc)
                
                if all_stage_docs:
                    # Используем эти документы для ответа
                    context = self._create_context_from_docs(all_stage_docs[:3])
                    docs = all_stage_docs[:3]
        
        try:
            # Упрощенный строгий промпт
            system_prompt = """Ты - AI Project Manager Assistant. Отвечай ТОЛЬКО на основе предоставленного контекста.
Используй информацию из контекста максимально полно.
Если информации в контексте достаточно для ответа - дай полный ответ.
Если информации недостаточно - скажи, что именно есть в контексте."""

            user_prompt = f"""Вопрос: {query}

Информация из базы знаний:
{context}

Ответь на вопрос используя ТОЛЬКО информацию из базы знаний выше.
Не добавляй свои знания.
Если в информации выше есть ответ - дай его.
Если нет - так и скажи.

Ответ:"""

            response = self.openai_client.chat.completions.create(
                model=self.main_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=800
            )
            
            answer = response.choices[0].message.content
            
            # Дополнительная проверка: если ответ слишком короткий, добавляем информацию из документов
            if len(answer.strip()) < 100 and docs:
                # Создаем дополнительный ответ на основе документов
                supplemental = "\n\n**Из документации:**\n"
                for i, doc in enumerate(docs[:2]):
                    # Берем первые 150 символов каждого документа
                    content = doc.page_content
                    if len(content) > 150:
                        content = content[:150] + "..."
                    supplemental += f"\n{i+1}. {content}"
                
                answer += supplemental
            
            return answer
            
        except Exception as e:
            print(f"⚠️ Ошибка генерации: {str(e)[:100]}")
            return f"Извините, произошла ошибка. Найдено документов: {len(docs) if docs else 0}"
    
    def _create_context_from_docs(self, docs: List[Document]) -> str:
        """Создание контекста из документов"""
        context_parts = []
        for i, doc in enumerate(docs):
            content = doc.page_content
            category = doc.metadata.get("category", "неизвестно")
            
            if len(content) > 500:
                content = content[:450] + "..."
            
            context_parts.append(f"Документ {i+1} ({category}):\n{content}")
        
        return "\n\n" + "="*50 + "\n\n".join(context_parts) + "\n" + "="*50
    
    
    def _no_info_response(self, query: str) -> str:
        """Ответ при отсутствии информации"""
        return f"❌ В базе знаний нет информации по запросу: '{query}'\n\n" \
               "Доступные темы в базе знаний:\n" \
               "1. 📋 **Этапы управления ML-проектом** (инициализация, декомпозиция, планирование и т.д.)\n" \
               "2. ⚠️ **Решение типовых проблем ML-проектов** (переобучение, дисбаланс классов и др.)\n" \
               "3. 🛠️ **Технологический стек** (PyTorch, TensorFlow, MLflow, Docker и др.)\n" \
               "4. ✅ **Шаблоны и чек-листы** (чек-лист старта проекта, чек-лист перед сдачей)\n" \
               "5. 📘 **Общие принципы** (миссия, подход, тон общения)\n\n" \
               "Попробуйте задать вопрос по одной из этих тем."
    
    def _extract_topic_summary(self, docs: List[Document]) -> str:
        """Извлечение краткого содержания найденных документов"""
        if not docs:
            return "Документы не найдены."
        
        summary_parts = []
        for i, doc in enumerate(docs[:3]):
            category = doc.metadata.get("category", "разное")
            content = doc.page_content
            
            # Извлекаем первые предложения
            sentences = content.split('.')
            brief = '.'.join(sentences[:2]) + '.' if len(sentences) > 2 else content[:150] + "..."
            
            summary_parts.append(f"{i+1}. **{category}**: {brief}")
        
        return "\n".join(summary_parts)
    
    def process_query(self, query: str) -> str:
        """Обработка запроса"""
        print(f"\n📥 Запрос: {query}")
        
        # Определяем категорию
        query_lower = query.lower()
        
        if any(word in query_lower for word in ["привет", "здравствуй", "hello", "hi"]):
            return self._greeting_response()
        
        # Определяем категорию по ключевым словам
        category_keywords = {
            "технологии": ["pytorch", "tensorflow", "mlflow", "docker", "фреймворк", "стек", "технологич"],
            "этапы": ["этап", "фаза", "шаг", "планирован", "инициализац", "дек", "позиц", 
                     "разбить", "подзадач", "декомпозиц"],
            "проблемы": ["проблем", "ошибк", "переобуч", "дисбаланс", "нехватк"],
            "чеклисты": ["чек-лист", "чеклист", "checklist", "шаблон"],
            "общее": ["миссия", "принцип", "цель", "подход", "общ"]
        }
        
        category = "другое"
        for cat, keywords in category_keywords.items():
            if any(keyword in query_lower for keyword in keywords):
                category = cat
                break
        
        print(f"   📍 Определена категория: '{category}'")
        
        # Поиск в базе знаний
        print("🔍 Этап 2: Поиск в базе знаний...")
        context, docs = self._retrieve_context(query, category)
        
        # Генерация ответа
        print(f"🧠 Этап 3: Генерация ответа ({self.main_model})...")
        answer = self._generate_strict_response(query, context, category, docs)
        
        print(f"   ✓ Ответ готов ({len(answer)} символов)")
        
        # Форматирование
        emoji_map = {
            "технологии": "🛠️",
            "этапы": "📋", 
            "проблемы": "⚠️",
            "чеклисты": "✅",
            "общее": "📘",
            "архитектура": "🏗️",
            "другое": "💡"
        }
        
        emoji = emoji_map.get(category, "🤖")
        formatted_response = f"{emoji} **Категория:** {category}\n\n{answer}"
        
        return formatted_response
    
    def _greeting_response(self) -> str:
        """Приветственное сообщение"""
        return "👋 Привет! Я AI-менеджер проектов для курса 'Разработчик нейросетей'.\n\n" \
               "Я могу ответить на вопросы о:\n\n" \
               "📋 **Этапах ML-проекта** (как разбить на подзадачи, планирование и т.д.)\n" \
               "⚠️ **Решение проблем** (переобучение, дисбаланс данных и др.)\n" \
               "🛠️ **Технологическом стеке** (PyTorch, MLflow, Docker и др.)\n" \
               "✅ **Чек-листах и шаблонах** (чек-лист старта, чек-лист сдачи)\n\n" \
               "Задайте вопрос по одной из этих тем!"

def create_gradio_interface(docx_path: str):
    """Создание Gradio интерфейса"""
    
    if os.path.exists("./chroma_neuro_db"):
        print("🗑️ Удаление старой векторной базы...")
        import shutil
        try:
            shutil.rmtree("./chroma_neuro_db")
            print("✅ Старая база удалена")
        except:
            print("⚠️ Не удалось удалить старую базу")
    
    try:
        manager = NeuroProjectManagerRAG(docx_path)
        print("🚀 Нейро-Менеджер Проектов успешно инициализирован!")
    except Exception as e:
        error_msg = f"❌ Ошибка инициализации: {e}"
        print(error_msg)
        
        class DummyManager:
            def process_query(self, query):
                return f"Система не инициализирована. Ошибка: {e}"
        
        manager = DummyManager()
    
    def respond(message: str, history: List[List[str]]):
        response = manager.process_query(message)
        yield response
    
    demo = gr.ChatInterface(
        fn=respond,
        title="🤖 Нейро-Менеджер Проектов (улучшенный поиск)",
        description="AI Assistant для управления ML-проектами\n" \
                   "✅ Улучшенный поиск релевантной информации\n" \
                   "❌ Не выдумывает ответы\n" \
                   "📋 Специальная обработка запросов о декомпозиции",
        examples=[
            "Как разбить проект на подзадачи?",
            "Какие этапы включает ML-проект?",
            "Что делать при переобучении модели?",
            "Какой технологический стек рекомендуется?"
        ]
    )
    
    return demo

def main():
    DOCX_FILE_PATH = "base.docx"
    
    if not os.path.exists(DOCX_FILE_PATH):
        print(f"❌ Файл не найден: {DOCX_FILE_PATH}")
        print(f"   Создайте файл base.docx с базой знаний")
        return
    
    print("=" * 60)
    print("🚀 ЗАПУСК НЕЙРО-МЕНЕДЖЕРА ПРОЕКТОВ (УЛУЧШЕННАЯ ВЕРСИЯ)")
    print("=" * 60)
    
    try:
        demo = create_gradio_interface(DOCX_FILE_PATH)
        
        print("\n🌐 Веб-интерфейс запускается...")
        print("📡 Откройте в браузере: http://localhost:7860")
        
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            share=False,
            show_error=True
        )
        
    except Exception as e:
        print(f"❌ Критическая ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()